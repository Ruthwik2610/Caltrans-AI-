import docx
import os
import sys

def extract_text(file_path):
    if not os.path.exists(file_path):
        return f"File {file_path} not found."
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"Error: {str(e)}"

if __name__ == "__main__":
    file_path = "CMGC Post Demo feedback requiremnts enhancement.docx"
    if os.path.exists(file_path):
        print(extract_text(file_path))
    else:
        print(f"Sample file {file_path} not found.")
