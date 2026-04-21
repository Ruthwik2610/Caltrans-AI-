import json
import datetime
import re
import logging
from io import BytesIO
from openai import OpenAI
from docx import Document
from src.delivery_method_kb import DELIVERY_METHOD_KB_TEXT

def _get_client(model_name: str = "gpt-4o"):
    """
    Universal LLM Client Switcher:
    - Defaults to standard OpenAI/Groq if API keys are found.
    - Falls back to Databricks Model Serving if running in a Databricks App.
    """
    import os
    from openai import OpenAI

    # 1. Check for standard Groq environment (Optimized for speed)
    if any(m in model_name.lower() for m in ["groq", "llama-3.3"]):
        api_key = os.getenv("GROQ_API_KEY")
        if api_key:
            return OpenAI(api_key=api_key, base_url="https://api.groq.com/openai/v1")

    # 2. Check for standard OpenAI environment
    if any(m in model_name.lower() for m in ["gpt", "openai"]):
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            return OpenAI(api_key=api_key)

    # 3. Fallback to Databricks (Local or Enterprise)
    try:
        from databricks.sdk import WorkspaceClient
        w = WorkspaceClient()
        return w.serving_endpoints.get_open_ai_client()
    except Exception:
        # Final fallback for local Databricks debugging
        token = os.getenv("DATABRICKS_TOKEN")
        host = os.getenv("DATABRICKS_HOST")
        if token and host:
            return OpenAI(api_key=token, base_url=f"{host.rstrip('/')}/serving-endpoints")
        
        # If all else fails, try standard OpenAI one last time with whatever is in env
        return OpenAI()

def _extract_json(text: str, finish_reason: str = "unknown") -> dict:
    """
    Robustly extract JSON from a string that might contain markdown blocks or leading/trailing text.
    Provides diagnostic info if parsing fails.
    """
    if not text or not text.strip():
        raise ValueError(f"AI response was empty (Finish Reason: {finish_reason}).")

    # Clean the input
    text = text.strip()

    # Try direct parse first (fastest)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try regex for markdown block ```json ... ```
    match = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Try regex for just ``` ... ```
    match = re.search(r"```\s*([\s\S]*?)```", text)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Last resort: find start and end braces
    first_brace = text.find("{")
    last_brace = text.rfind("}")
    if first_brace != -1 and last_brace != -1:
        try:
            return json.loads(text[first_brace:last_brace+1])
        except json.JSONDecodeError:
            pass

    # If all failed, provide a sample of the actual response for debugging
    sample = (text[:200] + "...") if len(text) > 200 else text
    raise ValueError(f"Could not parse AI response as JSON. Finish Reason: {finish_reason}. Raw Output Start: {sample}")

# ==============================================================================
# RUBRIC: 25 Questions extracted from AltDeliveryNominFactSheet Tables 3-8
# ==============================================================================
RUBRIC_QUESTIONS = [
    # --- Section A: Project Scope & Characteristics (Table 3, A1-A10) ---
    {
        "id": "A1", "section": "A: Project Scope & Characteristics",
        "question": "Where is the Project in the project development process?",
        "option_a": "Detailed or final engineering stage (60% design or later).",
        "option_b": "Preliminary design (30% design).",
        "option_c": "Conceptual engineering stage (before PA&ED).",
    },
    {
        "id": "A2", "section": "A: Project Scope & Characteristics",
        "question": "What is the size of the Project?",
        "option_a": "Small project (less than $25 million construction capital cost).",
        "option_b": "Medium size project (between $25 to $75 million construction capital cost).",
        "option_c": "Large project (greater than $75 million construction capital cost).",
    },
    {
        "id": "A3", "section": "A: Project Scope & Characteristics",
        "question": "What is the complexity of the Project?",
        "option_a": "Relatively simple project with no need for specialized outside expertise.",
        "option_b": "Project with more technically complex components and schedule complexity.",
        "option_c": "Very complex project with significant schedule complexity (e.g., multiple phases, extensive third-party issues, and/or specialized expertise needed).",
    },
    {
        "id": "A4", "section": "A: Project Scope & Characteristics",
        "question": "Does the Project involve significant impacts to highway users and local businesses/community during construction?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A5", "section": "A: Project Scope & Characteristics",
        "question": "Does the Project present right of way limitations that would benefit from the Entity's assistance?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A6", "section": "A: Project Scope & Characteristics",
        "question": "Does the Project present environmental permitting issues that would benefit from the Entity's assistance?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A7", "section": "A: Project Scope & Characteristics",
        "question": "Does the Project present utility or third-party issues that would benefit from the Entity's assistance?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A8", "section": "A: Project Scope & Characteristics",
        "question": "Does the Project present unique work restrictions (e.g., strict environmental windows, railroad restrictions) or traffic maintenance requirements that would benefit from the Entity's assistance?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A9", "section": "A: Project Scope & Characteristics",
        "question": "Would the Project benefit by packaging features of work to allow early lock-in of construction materials/labor pricing?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "A10", "section": "A: Project Scope & Characteristics",
        "question": "Would the Project benefit by raising quality standards/benchmarks to minimize maintenance and achieve lower life-cycle cost?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    # --- Section B: Schedule Issues (Table 4, B1-B2) ---
    {
        "id": "B1", "section": "B: Schedule Issues",
        "question": "Can time savings be realized through concurrent design and construction activities (fast-tracking)?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "B2", "section": "B: Schedule Issues",
        "question": "Can the schedule be compressed?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    # --- Section C: Opportunity for Innovation (Table 5, C1-C2) ---
    {
        "id": "C1", "section": "C: Opportunity for Innovation",
        "question": "Will the Project scope allow for innovation (e.g., alternate designs, traffic management, construction means and methods, etc.)?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "C2", "section": "C: Opportunity for Innovation",
        "question": "Must the Project scope be primarily defined in terms of prescriptive specifications, or can performance specifications be used, or a combination of both?",
        "option_a": "Primarily prescriptive specifications.",
        "option_b": "Combination of prescriptive and performance specifications.",
        "option_c": "Performance specifications for significant elements.",
    },
    # --- Section D: Quality Enhancement (Table 6, D1-D3) ---
    {
        "id": "D1", "section": "D: Quality Enhancement",
        "question": "Will there be opportunities for the Entity to provide materials or methods that provide greater value than normally specified by the state on similar projects?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "D2", "section": "D: Quality Enhancement",
        "question": "Will there be the opportunity for realization of greater value due to designs tailored to Entity's area of expertise?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "D3", "section": "D: Quality Enhancement",
        "question": "Will warranties or maintenance agreements be used?",
        "option_a": "No.",
        "option_b": "Limited to short-term workmanship and materials.",
        "option_c": "Much more than typical.",
    },
    # --- Section E: Cost Issues (Table 7, E1-E5) ---
    {
        "id": "E1", "section": "E: Cost Issues",
        "question": "Will there be opportunities for the Entity to provide designs with lower initial construction costs than those typically specified by the state?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "E2", "section": "E: Cost Issues",
        "question": "Will there be opportunities for the Entity to provide alternate design concepts with lower lifecycle costs than those typically specified by the state?",
        "option_a": "No more than typical.",
        "option_b": "More than typical.",
        "option_c": "Much more than typical.",
    },
    {
        "id": "E3", "section": "E: Cost Issues",
        "question": "Is funding for the Project committed and available?",
        "option_a": "Secured for design phase only or cannot support accelerated construction.",
        "option_b": "Funding can accommodate fast-tracking to some extent.",
        "option_c": "Funding will accommodate compressed schedule/fast-tracking.",
    },
    {
        "id": "E4", "section": "E: Cost Issues",
        "question": "Will the cost of procurement affect the number of bidders?",
        "option_a": "Procurement cost would significantly limit competition.",
        "option_b": "Procurement cost could affect the number of bidders.",
        "option_c": "Procurement cost would not be a significant issue given the size or complexity of the Project.",
    },
    {
        "id": "E5", "section": "E: Cost Issues",
        "question": "Will Project budget control benefit from the use of formal contingencies?",
        "option_a": "No benefit.",
        "option_b": "A formal contingency may permit the Department to add Project scope or enhance quality within the constraints of its published budget.",
        "option_c": "A formal contingency is required to allow the Department to maximize Project scope and quality within the constraints of its published budget.",
    },
    # --- Section F: Staffing Issues (Table 8, F1-F3) ---
    {
        "id": "F1", "section": "F: Staffing Issues",
        "question": "Does the Department have the expertise and resources necessary for a complicated procurement process?",
        "option_a": "Inadequate resources or expertise.",
        "option_b": "Limited resources or expertise.",
        "option_c": "Adequate resources and expertise.",
    },
    {
        "id": "F2", "section": "F: Staffing Issues",
        "question": "Are resources available to complete the design?",
        "option_a": "Resources are available to complete design.",
        "option_b": "Resources are available for partial design.",
        "option_c": "Specialized expertise, not available in-house, is required.",
    },
    {
        "id": "F3", "section": "F: Staffing Issues",
        "question": "Are resources available to provide construction oversight?",
        "option_a": "Resources are available.",
        "option_b": "Full-time construction oversight could strain staff resources.",
        "option_c": "Resources are unavailable.",
    },
]

# Section weights for scoring matrix
SECTION_WEIGHTS = {
    "A": 0.30,  # Project Scope & Characteristics (10 questions)
    "B": 0.15,  # Schedule Issues (2 questions)
    "C": 0.12,  # Opportunity for Innovation (2 questions)
    "D": 0.10,  # Quality Enhancement (3 questions)
    "E": 0.20,  # Cost Issues (5 questions)
    "F": 0.13,  # Staffing Issues (3 questions)
}

RATING_VALUES = {"A": 1, "B": 2, "C": 3}


# ==============================================================================
# DOCUMENT EXTRACTION
# ==============================================================================
def extract_text_from_uploaded_pdf(file) -> str:
    """Extract text from an uploaded PDF file object."""
    from PyPDF2 import PdfReader
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def extract_multi_doc_context(files) -> str:
    """Extract text from multiple files (DOCX or PDF) and wrap them in XML tags."""
    combined_text = []
    for file in files:
        name = getattr(file, "name", "Unknown Document")
        try:
            if name.lower().endswith(".docx"):
                doc = Document(file)
                content = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            elif name.lower().endswith(".pdf"):
                content = extract_text_from_uploaded_pdf(file)
            else:
                content = "UNSUPPORTED FILE TYPE"
            combined_text.append(f'<source_document name="{name}">\n{content}\n</source_document>')
        except Exception as e:
            combined_text.append(f'<source_document name="{name}">\nERROR EXTRACTING: {str(e)}\n</source_document>')
    return "\n\n".join(combined_text)


def extract_text_from_docx(file) -> tuple:
    """Extract narrative text and any pre-filled ratings from a nomination fact sheet DOCX.

    Returns:
        (narrative_text, existing_ratings) where existing_ratings is a dict like {"A1": "B", ...}
    """
    doc = Document(file)

    # Extract paragraph text (Sections 1-12, stop at Section 13)
    narrative_parts = []
    stop_keywords = [
        "Project Risk Assessment",
        "Construction Manager Tasks",
        "Glossary of Preconstruction",
        "District Single Point Signature",
    ]
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(kw.lower() in text.lower() for kw in stop_keywords):
            break
        narrative_parts.append(text)
    narrative_text = "\n".join(narrative_parts)

    # Extract pre-filled ratings from Tables 3-8 (evaluation question tables)
    existing_ratings = {}
    for table in doc.tables:
        if len(table.columns) != 3:
            continue
        header_cells = [cell.text.strip().upper() for cell in table.rows[0].cells]
        if "QUESTION NO." not in header_cells[0] and "QUESTION" not in header_cells[0]:
            continue
        # This is a rubric question table
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            question_id = cells[0].strip()
            rating_text = cells[2].strip().upper() if len(cells) > 2 else ""
            if question_id and rating_text in ("A", "B", "C"):
                existing_ratings[question_id] = rating_text

    return narrative_text, existing_ratings


def load_delivery_method_kb() -> str:
    """Return the embedded delivery method knowledge base text."""
    return DELIVERY_METHOD_KB_TEXT


# ==============================================================================
# SYSTEM PROMPT CONSTRUCTION
# ==============================================================================
def _build_rubric_text() -> str:
    """Format the 25 rubric questions for prompt injection."""
    lines = ["RUBRIC - 25 EVALUATION QUESTIONS:", "For each question, select exactly one rating: A (first option), B (second option), or C (third option).", ""]
    current_section = ""
    for q in RUBRIC_QUESTIONS:
        if q["section"] != current_section:
            current_section = q["section"]
            lines.append(f"--- {current_section} ---")
        lines.append(f'[{q["id"]}] {q["question"]}')
        lines.append(f'  A: {q["option_a"]}')
        lines.append(f'  B: {q["option_b"]}')
        lines.append(f'  C: {q["option_c"]}')
        lines.append("")
    return "\n".join(lines)


def _build_system_prompt(kb_text: str, existing_ratings: dict = None) -> str:
    """Build the full system prompt for the delivery evaluation."""

    persona = """You are a Senior Alternative Contracting Expert at Caltrans Headquarters, Office of Innovative Design and Delivery (OIDD). You have 20+ years of experience evaluating project delivery method nominations across California. Your role is to objectively evaluate a district's nomination fact sheet against the 25-question delivery selection rubric.

You are meticulous, evidence-based, and transparent about uncertainty. When the narrative lacks information for a question, you flag it clearly rather than guessing confidently."""

    kb_section = f"""DELIVERY METHOD COMPARISON KNOWLEDGE BASE:
Use the following reference to understand how each delivery method performs across different project factors (Project Requirements, Delivery Schedule, Complexity & Innovation, Level of Design, Cost, Risk Characteristics, Site Conditions, Utilities, Environmental, ROW, Third-Party Involvement):

{kb_text}"""

    design_sequencing = """ADDITIONAL METHOD - DESIGN-SEQUENCING:
Design-Sequencing is NOT in the comparison PDF above. Key characteristics:
- A variant of Design-Bid-Build where design packages are released sequentially
- Allows construction to begin on early packages while later packages are still being designed
- Department retains full design control (like DBB)
- Lower procurement complexity than DB or CMGC
- Appropriate when: project is moderately complex, schedule benefits from partial overlap, but full design-build risk transfer is not warranted
- Does not require contractor input during design (unlike CMGC)"""

    rubric_text = _build_rubric_text()

    baseline_norms = """CALTRANS BASELINE NORMS (apply these quantitative thresholds when data is available):
- A1 (Development Stage): 60%+ design = A, ~30% design = B, before PA&ED = C
- A2 (Project Size): <$25M construction capital = A, $25-75M = B, >$75M = C
- E3 (Funding): Secured for design only = A, can accommodate some fast-tracking = B, full compressed schedule = C
- E4 (Procurement Cost): Significantly limits competition = A, could affect bidders = B, not significant = C
- F1 (Procurement Expertise): Inadequate = A, Limited = B, Adequate = C
- F2 (Design Resources): Available to complete = A, Available for partial = B, Specialized expertise needed = C
- F3 (Construction Oversight): Available = A, Could strain staff = B, Unavailable = C
For questions using "No more than typical / More than typical / Much more than typical", use these guidelines:
- A (No more than typical): Standard Caltrans project, no unusual factors
- B (More than typical): Notable complexity or needs beyond standard, but manageable
- C (Much more than typical): Exceptional complexity, significant challenges requiring specialized approaches"""

    few_shot = """EVALUATION METHODOLOGY:
For each of the 25 questions, follow this chain-of-thought:
1. EXTRACT: Find all evidence in the narrative relevant to this question. Quote or summarize directly.
2. ANALYZE: Apply the rubric criteria. Compare evidence against options A, B, and C.
3. RATE: Select the rating that best matches.
4. FLAG: If insufficient evidence, set missing_info to true and estimate using domain knowledge. Lower confidence to 0.2-0.4 for estimates.

When a question has no direct evidence, check if related questions provide indirect evidence. For example, if A3 (complexity) has evidence of high complexity, that indirectly supports higher ratings for C1 (innovation opportunity) and B1 (fast-tracking benefit).

EXAMPLE 1 - Question A2 (Project Size):
source_reasoning: "The construction capital cost is estimated at $45 million."
missing_info_reasoning: "Funding is confirmed for construction, but no specific utility relocation estimates are provided yet."
effect_on_method: "$45M falls within the standard Design-Build range. The confirmed funding favors Design-Build over traditional DBB for better cost control."
selected_rating: "B"
confidence: 0.95
missing_info: false

EXAMPLE 2 - Question A7 (Utility/Third-Party Issues):
source_reasoning: "The project requires coordination with BNSF railroad for track closures and PG&E for gas line relocation. Multiple utility relocations are anticipated."
missing_info_reasoning: "Specific right-of-way entry agreements with BNSF are not yet in place."
effect_on_method: "Railroad and utility coordination involving multiple parties goes beyond typical. CMGC or PDB would be beneficial to bring in a contractor early to manage these risks."
selected_rating: "B"
confidence: 0.85
missing_info: false

EXAMPLE 3 - Question C1 (Innovation) with missing info:
source_reasoning: "No explicit discussion of innovation opportunities found in the narrative."
missing_info_reasoning: "The narrative lacks a section on innovation potential or alternative technical concepts."
effect_on_method: "The narrative does not address innovation. Based on the project's complexity (if A3 suggests moderate complexity) and highway scope, moderate innovation potential is plausible but unconfirmed."
selected_rating: "B"
confidence: 0.35
missing_info: true"""

    exclusion = """IMPORTANT EXCLUSIONS:
- Do NOT evaluate any content from Sections 13, 14, or 15 of the fact sheet (Risk Register, CMGC Task Selection, Glossary).
- Evaluate ONLY based on the project narrative from Sections 1-12 and any supplementary project details.
- Your evaluation must cover ALL 25 questions (A1 through F3). Do not skip any."""

    existing_ratings_text = ""
    if existing_ratings:
        ratings_str = ", ".join(f"{k}={v}" for k, v in sorted(existing_ratings.items()))
        existing_ratings_text = f"""
DISTRICT PRE-FILLED RATINGS:
The district has pre-filled these ratings: {ratings_str}
Evaluate independently based on the evidence. After your independent evaluation, if your rating differs from the district's, note the disagreement in your effect_on_method."""

    output_schema = """OUTPUT FORMAT:
You must output ONLY valid JSON in the following format. Replace all placeholders with real values extracted from the narrative.

{
  "project_name": "<extract project name or description from the narrative>",
  "project_ea": "<extract Project EA number or NOT PROVIDED>",
  "district": "<extract district name/number or NOT PROVIDED>",
  "evaluation_date": "<today's date in YYYY-MM-DD format>",
  "ratings": [
    {
      "question_id": "A1",
      "question_text": "Where is the Project in the project development process?",
      "source_reasoning": "<direct quote or summary from narrative, or 'No direct evidence found'>",
      "missing_info_reasoning": "<explicit explanation of what information is missing or 'None'>",
      "effect_on_method": "<analysis of how this rating affects the suitability of the delivery method and comparison to rubric norms>",
      "selected_rating": "A or B or C",
      "confidence": 0.0 to 1.0,
      "missing_info": true or false
    }
  ],
  "missing_questions": ["list of question_ids where missing_info is true"],
  "summary": "<2-3 sentence overall assessment of the project and evaluation quality>"
}

CRITICAL: The "ratings" array must contain EXACTLY 25 items, one for each question A1 through F3, in order."""

    return "\n\n".join([
        persona, kb_section, design_sequencing, rubric_text,
        baseline_norms, few_shot, exclusion, existing_ratings_text, output_schema
    ])


# ==============================================================================
# CORE EVALUATION FUNCTION
# ==============================================================================
def run_delivery_evaluation(narrative_text: str, kb_text: str, existing_ratings: dict = None, model_name: str = "gpt-4o") -> dict:
    """Run the 25-question delivery method evaluation using an LLM.

    Args:
        narrative_text: Extracted text from the nomination fact sheet
        kb_text: Text from the DeliveryMethodComparison PDF
        existing_ratings: Optional dict of district pre-filled ratings {"A1": "B", ...}
        model_name: The LLM model/endpoint to use (defaults to gpt-4o)

    Returns:
        Dict with evaluation results or {"error": "..."} on failure
    """
    system_prompt = _build_system_prompt(kb_text, existing_ratings)

    user_message = f"""Please evaluate the following Alternative Delivery Nomination Fact Sheet against all 25 rubric questions.

NOMINATION FACT SHEET CONTENT:
{narrative_text}"""

    try:
        client = _get_client(model_name)
        response = client.chat.completions.create(
            model=model_name,
            response_format={"type": "json_object"} if any(m in model_name.lower() for m in ["gpt", "json", "gemma"]) else None,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.0,
        )
        
        main_text = response.choices[0].message.content
        main_reason = response.choices[0].finish_reason
        
        try:
            return _extract_json(main_text, finish_reason=main_reason)
        except (json.JSONDecodeError, ValueError):
            # Retry once on malformed JSON
            retry = client.chat.completions.create(
                model=model_name,
                response_format={"type": "json_object"} if any(m in model_name.lower() for m in ["gpt", "json", "gemma"]) else None,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message},
                ],
                temperature=0.0,
            )
            retry_text = retry.choices[0].message.content
            retry_reason = retry.choices[0].finish_reason
            return _extract_json(retry_text, finish_reason=retry_reason)
    except Exception as e:
        return {"error": f"AI service error during delivery evaluation: {str(e)}"}


# ==============================================================================
# SCORING MATRIX & DELIVERY METHOD RECOMMENDATION
# ==============================================================================
def compute_delivery_recommendation(ratings: list) -> dict:
    """Apply weighted scoring matrix to 25 A/B/C ratings and recommend a delivery method.

    Args:
        ratings: List of 25 dicts with "question_id" and "selected_rating" keys

    Returns:
        Dict with composite_score, section_scores, recommended_method, etc.
    """
    # Group ratings by section
    section_ratings = {"A": [], "B": [], "C": [], "D": [], "E": [], "F": []}
    rating_lookup = {}
    for r in ratings:
        qid = r.get("question_id", "")
        rating = r.get("selected_rating", "B").upper()
        if qid and qid[0] in section_ratings:
            section_ratings[qid[0]].append(RATING_VALUES.get(rating, 2))
            rating_lookup[qid] = rating

    # Compute section averages
    section_scores = {}
    for section, values in section_ratings.items():
        section_scores[section] = sum(values) / len(values) if values else 2.0

    # Compute weighted composite score
    composite = sum(
        section_scores[s] * SECTION_WEIGHTS[s] for s in SECTION_WEIGHTS
    )

    # Determine recommended method
    recommended, runner_up, is_borderline, comparison = _determine_method(
        composite, section_scores, rating_lookup
    )

    # Apply document-based override rules
    recommended, runner_up, override_reasons = _apply_overrides(
        recommended, runner_up, rating_lookup
    )

    # Recompute borderline comparison if override changed the recommendation
    if override_reasons and is_borderline:
        comparison = _build_comparison(recommended, runner_up, composite, section_scores)

    # --- Phase 2 Enhancements ---
    # Raw scores: numeric value per question
    raw_scores = {qid: RATING_VALUES.get(r, 2) for qid, r in rating_lookup.items()}

    # Weighted scores: section_avg × section_weight
    weighted_scores = {s: round(section_scores[s] * SECTION_WEIGHTS[s], 4) for s in SECTION_WEIGHTS}

    # Override status: evaluate all 9 rules, mark triggered or not
    override_status = _compute_override_status(rating_lookup)

    # Key drivers: top 5 questions by weighted contribution
    question_weights = []
    for qid, raw in raw_scores.items():
        sec = qid[0]
        sec_count = len(section_ratings.get(sec, [1]))
        per_q_weight = SECTION_WEIGHTS.get(sec, 0) / max(sec_count, 1)
        question_weights.append({
            "question_id": qid,
            "raw_score": raw,
            "rating": rating_lookup.get(qid, "B"),
            "section": sec,
            "weighted_contribution": round(raw * per_q_weight, 4),
        })
    question_weights.sort(key=lambda x: x["weighted_contribution"], reverse=True)
    key_drivers = question_weights[:5]

    return {
        "composite_score": round(composite, 3),
        "section_scores": {k: round(v, 3) for k, v in section_scores.items()},
        "raw_scores": raw_scores,
        "weighted_scores": weighted_scores,
        "override_status": override_status,
        "key_drivers": key_drivers,
        "recommended_method": recommended,
        "runner_up_method": runner_up,
        "is_borderline": is_borderline,
        "comparison_text": comparison,
        "override_reasons": override_reasons,
    }


# ==============================================================================
# MULTI-METHOD SCORING  (Req 3.1)
# ==============================================================================

# Method-affinity matrix derived from the Caltrans Delivery Method Comparison KB.
# Each section maps to how strongly an "A" vs "C" rating favors each method.
# Values: 1.0 = strongly favors, 0.5 = neutral, 0.0 = disfavors.
# A rating of "C" (high complexity/need) is mapped using the "C" column;
# "A" (simple/traditional) uses the "A" column; "B" interpolates.
METHOD_AFFINITY = {
    # Section A: Project Scope & Characteristics — high complexity favors CMGC/DB/PDB
    "A": {
        "Design-Bid-Build":        {"A": 1.0, "B": 0.6, "C": 0.1},
        "Design-Sequencing":       {"A": 0.8, "B": 0.6, "C": 0.2},
        "Design-Build/Low-Bid":    {"A": 0.3, "B": 0.5, "C": 0.7},
        "Design-Build/Best-Value": {"A": 0.2, "B": 0.5, "C": 0.8},
        "CM/GC":                   {"A": 0.3, "B": 0.6, "C": 0.9},
        "Progressive Design-Build":{"A": 0.1, "B": 0.4, "C": 0.9},
    },
    # Section B: Schedule Issues — urgency favors DB, CMGC, PDB
    "B": {
        "Design-Bid-Build":        {"A": 0.9, "B": 0.5, "C": 0.1},
        "Design-Sequencing":       {"A": 0.7, "B": 0.5, "C": 0.3},
        "Design-Build/Low-Bid":    {"A": 0.3, "B": 0.6, "C": 0.8},
        "Design-Build/Best-Value": {"A": 0.3, "B": 0.6, "C": 0.8},
        "CM/GC":                   {"A": 0.4, "B": 0.6, "C": 0.8},
        "Progressive Design-Build":{"A": 0.3, "B": 0.5, "C": 0.8},
    },
    # Section C: Innovation — high innovation favors DB/BV, CMGC, PDB
    "C": {
        "Design-Bid-Build":        {"A": 0.9, "B": 0.5, "C": 0.1},
        "Design-Sequencing":       {"A": 0.7, "B": 0.5, "C": 0.2},
        "Design-Build/Low-Bid":    {"A": 0.4, "B": 0.5, "C": 0.6},
        "Design-Build/Best-Value": {"A": 0.2, "B": 0.5, "C": 0.9},
        "CM/GC":                   {"A": 0.3, "B": 0.6, "C": 0.8},
        "Progressive Design-Build":{"A": 0.2, "B": 0.5, "C": 0.9},
    },
    # Section D: Quality Enhancement — high quality needs favor DB/BV, CMGC
    "D": {
        "Design-Bid-Build":        {"A": 0.8, "B": 0.5, "C": 0.2},
        "Design-Sequencing":       {"A": 0.7, "B": 0.5, "C": 0.3},
        "Design-Build/Low-Bid":    {"A": 0.5, "B": 0.5, "C": 0.5},
        "Design-Build/Best-Value": {"A": 0.3, "B": 0.5, "C": 0.8},
        "CM/GC":                   {"A": 0.3, "B": 0.6, "C": 0.8},
        "Progressive Design-Build":{"A": 0.3, "B": 0.5, "C": 0.8},
    },
    # Section E: Cost Issues — constrained funding favors DBB; full funding favors DB
    "E": {
        "Design-Bid-Build":        {"A": 0.8, "B": 0.6, "C": 0.3},
        "Design-Sequencing":       {"A": 0.7, "B": 0.5, "C": 0.3},
        "Design-Build/Low-Bid":    {"A": 0.3, "B": 0.5, "C": 0.7},
        "Design-Build/Best-Value": {"A": 0.2, "B": 0.5, "C": 0.7},
        "CM/GC":                   {"A": 0.4, "B": 0.6, "C": 0.7},
        "Progressive Design-Build":{"A": 0.2, "B": 0.5, "C": 0.8},
    },
    # Section F: Staffing — lack of expertise favors DB, PDB, CMGC
    "F": {
        "Design-Bid-Build":        {"A": 0.8, "B": 0.5, "C": 0.1},
        "Design-Sequencing":       {"A": 0.7, "B": 0.5, "C": 0.2},
        "Design-Build/Low-Bid":    {"A": 0.4, "B": 0.5, "C": 0.6},
        "Design-Build/Best-Value": {"A": 0.3, "B": 0.5, "C": 0.7},
        "CM/GC":                   {"A": 0.5, "B": 0.6, "C": 0.7},
        "Progressive Design-Build":{"A": 0.3, "B": 0.5, "C": 0.8},
    },
}

ALL_METHODS = [
    "Design-Bid-Build", "Design-Sequencing", "Design-Build/Low-Bid",
    "Design-Build/Best-Value", "CM/GC", "Progressive Design-Build",
]

# Method pros/cons derived from the KB
_METHOD_PROS_CONS = {
    "Design-Bid-Build": {
        "pros": ["Lowest procurement cost", "Well-understood process", "Competitive bidding on fully defined scope", "Department retains full design control"],
        "cons": ["Longest delivery schedule", "No early work packages", "Higher change order risk", "All design risks on Department"],
    },
    "Design-Sequencing": {
        "pros": ["Sequential design packages allow phased delivery", "Department retains design control", "Lower procurement complexity than DB"],
        "cons": ["Still relatively slow", "Limited contractor input during design", "No three-party collaboration"],
    },
    "Design-Build/Low-Bid": {
        "pros": ["Parallel design and construction", "Risk transfer to Design-Builder", "Schedule compression possible"],
        "cons": ["Less focus on quality/innovation", "Higher procurement costs", "30% design needed for RFP", "Less Department control over final design"],
    },
    "Design-Build/Best-Value": {
        "pros": ["ATCs drive innovation", "Best quality-price balance", "Parallel design and construction", "Risk transfer to Design-Builder"],
        "cons": ["Highest procurement cost (stipends)", "Complex two-phase procurement", "Requires well-defined contract requirements"],
    },
    "CM/GC": {
        "pros": ["Three-party collaboration", "Early contractor input", "Early work packages possible", "Flexible agreed-price negotiation", "Department retains design ownership"],
        "cons": ["If agreed price fails, reverts to DBB", "Requires sophisticated procurement", "Preconstruction services add cost"],
    },
    "Progressive Design-Build": {
        "pros": ["Three-party collaboration", "Can start at PA&ED phase", "Maximum innovation opportunity", "Early work packages", "Design-Builder owns final design"],
        "cons": ["Requires >$25M project cost", "Complex procurement", "If price fails, must use other methods", "Department has less design control"],
    },
}


def score_all_methods(ratings: list) -> dict:
    """Score ALL 6 delivery methods using the method-affinity matrix.

    Returns dict with:
      - method_scores: list of {method, score, rank, blocked, pros, cons, key_factors}
      - borderline_comparison: detailed comparison if top 2 are close
    """
    # Build rating lookup
    rating_lookup = {}
    section_ratings = {"A": [], "B": [], "C": [], "D": [], "E": [], "F": []}
    for r in ratings:
        qid = r.get("question_id", "")
        rating = r.get("selected_rating", "B").upper()
        if qid and qid[0] in section_ratings:
            rating_lookup[qid] = rating
            section_ratings[qid[0]].append(rating)

    # Compute section-level dominant ratings (mode of A/B/C per section)
    section_dominant = {}
    for sec, rs in section_ratings.items():
        if not rs:
            section_dominant[sec] = "B"
        else:
            # Use average: A=3, B=2, C=1 -> map back
            avg = sum(RATING_VALUES.get(r, 2) for r in rs) / len(rs)
            if avg >= 2.5:
                section_dominant[sec] = "A"
            elif avg >= 1.5:
                section_dominant[sec] = "B"
            else:
                section_dominant[sec] = "C"

    # Score each method
    method_scores = []
    for method in ALL_METHODS:
        weighted_sum = 0.0
        key_factors = []
        for sec, weight in SECTION_WEIGHTS.items():
            dominant = section_dominant.get(sec, "B")
            affinity = METHOD_AFFINITY.get(sec, {}).get(method, {}).get(dominant, 0.5)
            contribution = affinity * weight
            weighted_sum += contribution
            if affinity >= 0.8:
                key_factors.append(f"{sec}: Strong fit ({dominant})")
            elif affinity <= 0.2:
                key_factors.append(f"{sec}: Poor fit ({dominant})")

        method_scores.append({
            "method": method,
            "score": round(weighted_sum, 4),
            "blocked": False,
            "block_reasons": [],
            "pros": _METHOD_PROS_CONS.get(method, {}).get("pros", []),
            "cons": _METHOD_PROS_CONS.get(method, {}).get("cons", []),
            "key_factors": key_factors[:5],
        })

    # Apply override blocks
    override_status = _compute_override_status(rating_lookup)
    blocked_methods = set()
    for o in override_status:
        if o["triggered"]:
            for b in o.get("blocks", []):
                blocked_methods.add(b)

    for ms in method_scores:
        if ms["method"] in blocked_methods:
            ms["blocked"] = True
            ms["block_reasons"] = [
                o["rule_id"] + ": " + o["rule_name"]
                for o in override_status
                if o["triggered"] and ms["method"] in o.get("blocks", [])
            ]

    # Rank (unblocked first, then by score descending)
    method_scores.sort(key=lambda x: (-int(not x["blocked"]), -x["score"]))
    for i, ms in enumerate(method_scores):
        ms["rank"] = i + 1

    # Borderline comparison for top 2-3
    borderline_comparison = None
    unblocked = [ms for ms in method_scores if not ms["blocked"]]
    if len(unblocked) >= 2 and (unblocked[0]["score"] - unblocked[1]["score"]) <= 0.05:
        top_methods = unblocked[:3] if len(unblocked) >= 3 and (unblocked[0]["score"] - unblocked[2]["score"]) <= 0.10 else unblocked[:2]
        borderline_comparison = {
            "is_close": True,
            "score_gap": round(unblocked[0]["score"] - unblocked[1]["score"], 4),
            "methods": [
                {
                    "method": m["method"],
                    "score": m["score"],
                    "pros": m["pros"],
                    "cons": m["cons"],
                    "key_factors": m["key_factors"],
                }
                for m in top_methods
            ],
        }

    return {
        "method_scores": method_scores,
        "borderline_comparison": borderline_comparison,
        "override_status": override_status,
    }


# ==============================================================================
# VALIDATION MODE  (Req 3.3 + 3.7)
# ==============================================================================
def run_validation_analysis(ai_ratings: list, user_ratings: dict) -> dict:
    """Compare AI ratings against user (district) ratings.

    Args:
        ai_ratings: list of rating dicts from run_delivery_evaluation
        user_ratings: dict of {question_id: rating} from the uploaded fact sheet

    Returns:
        Dict with comparison details, mismatch analysis, and deviation impact.
    """
    comparisons = []
    mismatches = []

    for r in ai_ratings:
        qid = r.get("question_id", "")
        ai_rating = r.get("selected_rating", "B").upper()
        user_rating = user_ratings.get(qid, "").upper()

        if not user_rating:
            continue

        ai_val = RATING_VALUES.get(ai_rating, 2)
        user_val = RATING_VALUES.get(user_rating, 2)
        diff = abs(ai_val - user_val)

        if diff == 0:
            severity = "match"
        elif diff == 1:
            severity = "minor_mismatch"
        else:
            severity = "major_mismatch"

        entry = {
            "question_id": qid,
            "question_text": r.get("question_text", ""),
            "ai_rating": ai_rating,
            "user_rating": user_rating,
            "severity": severity,
            "ai_evidence": r.get("extracted_evidence", ""),
            "ai_confidence": round(r.get("confidence", 0), 2),
            "has_evidence": r.get("confidence", 0) >= 0.4,
        }
        comparisons.append(entry)
        if severity != "match":
            mismatches.append(entry)

    total = len(comparisons)
    matches = sum(1 for c in comparisons if c["severity"] == "match")
    minor = sum(1 for c in comparisons if c["severity"] == "minor_mismatch")
    major = sum(1 for c in comparisons if c["severity"] == "major_mismatch")

    # Compute deviation impact: what would change if we use user ratings instead
    ai_rec = compute_delivery_recommendation(ai_ratings)
    user_modified_ratings = []
    for r in ai_ratings:
        qid = r.get("question_id", "")
        modified = dict(r)
        if qid in user_ratings and user_ratings[qid]:
            modified["selected_rating"] = user_ratings[qid]
        user_modified_ratings.append(modified)
    user_rec = compute_delivery_recommendation(user_modified_ratings)

    recommendation_changed = ai_rec["recommended_method"] != user_rec["recommended_method"]

    return {
        "comparisons": comparisons,
        "mismatches": mismatches,
        "summary": {
            "total_compared": total,
            "matches": matches,
            "minor_mismatches": minor,
            "major_mismatches": major,
            "agreement_rate": round(matches / max(total, 1) * 100, 1),
        },
        "deviation_impact": {
            "ai_method": ai_rec["recommended_method"],
            "user_method": user_rec["recommended_method"],
            "recommendation_changed": recommendation_changed,
            "ai_score": ai_rec["composite_score"],
            "user_score": user_rec["composite_score"],
        },
    }


def _determine_method(composite: float, section_scores: dict, rating_lookup: dict) -> tuple:
    """Determine the delivery method from composite and sub-scores.

    Returns: (recommended, runner_up, is_borderline, comparison_text)
    """
    thresholds = [
        (1.40, "Design-Bid-Build"),
        (1.70, "Design-Sequencing"),
        (2.10, None),  # Sub-score dependent
        (2.50, None),  # Sub-score dependent
        (3.01, "Progressive Design-Build"),
    ]

    # Check borderline (within 0.15 of a threshold boundary)
    is_borderline = False
    for t, _ in thresholds:
        if abs(composite - t) < 0.15:
            is_borderline = True
            break

    # Primary mapping
    if composite <= 1.40:
        recommended = "Design-Bid-Build"
        runner_up = "Design-Sequencing"
    elif composite <= 1.70:
        recommended = "Design-Sequencing"
        runner_up = "Design-Bid-Build" if composite < 1.55 else "CM/GC"
    elif composite <= 2.10:
        recommended, runner_up = _mid_range_method(section_scores, rating_lookup)
    elif composite <= 2.50:
        recommended, runner_up = _upper_range_method(section_scores, rating_lookup)
    else:
        recommended = "Progressive Design-Build"
        runner_up = "CM/GC"

    comparison = ""
    if is_borderline:
        comparison = _build_comparison(recommended, runner_up, composite, section_scores)

    return recommended, runner_up, is_borderline, comparison


def _mid_range_method(section_scores: dict, rating_lookup: dict) -> tuple:
    """Determine method for composite scores 1.71-2.10.

    In this mid-range, CM/GC is the most common recommendation because the
    project is complex enough to benefit from collaboration but not so
    extreme as to require full design-build risk transfer.
    """
    a_avg = section_scores.get("A", 2.0)
    b_avg = section_scores.get("B", 2.0)
    c_avg = section_scores.get("C", 2.0)
    d_avg = section_scores.get("D", 2.0)
    f_avg = section_scores.get("F", 2.0)

    if b_avg >= 2.5 and c_avg >= 2.0:
        # Strong schedule + innovation signals -> Design-Build
        return "Design-Build/Best-Value", "CM/GC"
    if b_avg >= 2.5 and c_avg < 2.0:
        return "Design-Build/Low-Bid", "CM/GC"
    # PDB only if complexity AND scope are both very high in mid-range
    if a_avg >= 2.5 and rating_lookup.get("A3") == "C" and rating_lookup.get("E3") == "C":
        return "Progressive Design-Build", "CM/GC"
    # CM/GC is the default for mid-range — collaborative approach for complex projects
    if c_avg >= 2.0:
        return "CM/GC", "Design-Build/Best-Value"
    return "CM/GC", "Design-Sequencing"


def _upper_range_method(section_scores: dict, rating_lookup: dict) -> tuple:
    """Determine method for composite scores 2.11-2.50."""
    c_avg = section_scores.get("C", 2.0)
    d_avg = section_scores.get("D", 2.0)
    f_avg = section_scores.get("F", 2.0)

    if c_avg < 1.5:
        return "Design-Build/Low-Bid", "CM/GC"
    if c_avg >= 2.0 and d_avg >= 2.0:
        return "Design-Build/Best-Value", "Progressive Design-Build"
    if f_avg >= 2.5:
        return "CM/GC", "Progressive Design-Build"
    if rating_lookup.get("A3") == "C" and rating_lookup.get("E3") == "C":
        return "Progressive Design-Build", "Design-Build/Best-Value"
    return "CM/GC", "Design-Build/Best-Value"


# Fallback hierarchy: most flexible → most constrained
_FALLBACK_ORDER = [
    "CM/GC",
    "Design-Build/Best-Value",
    "Design-Build/Low-Bid",
    "Progressive Design-Build",
    "Design-Sequencing",
    "Design-Bid-Build",
]

# Override rules with human-readable descriptions for UI display
OVERRIDE_RULES = [
    {
        "id": "R1",
        "name": "Early-stage design requires flexible procurement",
        "trigger": "A1 = C",
        "description": "Project is in conceptual stage (before PA&ED). Design-Bid-Build and Design-Sequencing both require advanced design completion before procurement.",
        "blocks": {"Design-Bid-Build", "Design-Sequencing"},
    },
    {
        "id": "R2",
        "name": "Large projects exceed DBB scope",
        "trigger": "A2 = C",
        "description": "Project exceeds $75M construction capital. Projects of this scale typically require collaborative or design-build delivery methods.",
        "blocks": {"Design-Bid-Build"},
    },
    {
        "id": "R3",
        "name": "Very complex projects need collaborative delivery",
        "trigger": "A3 = C",
        "description": "Very complex project with significant schedule complexity. Requires three-party collaboration (Department, Contractor, ICE) that DBB and Design-Sequencing do not provide.",
        "blocks": {"Design-Bid-Build", "Design-Sequencing"},
    },
    {
        "id": "R4",
        "name": "Maximum schedule compression blocks DBB",
        "trigger": "B1 = C AND B2 = C",
        "description": "Project requires maximum fast-tracking and schedule compression. DBB has the longest delivery schedule and does not allow early work packages.",
        "blocks": {"Design-Bid-Build"},
    },
    {
        "id": "R5",
        "name": "Performance specifications favor Design-Build",
        "trigger": "C2 = C",
        "description": "Project uses performance specifications for significant elements. Design-Build methods leverage ATCs (Alternative Technical Concepts) to meet performance outcomes.",
        "blocks": set(),
        "favor": "Design-Build/Best-Value",
    },
    {
        "id": "R6",
        "name": "Limited funding blocks accelerated methods",
        "trigger": "E3 = A",
        "description": "Funding is secured for design phase only. Design-Build and Progressive Design-Build require construction capital commitment for procurement.",
        "blocks": {"Design-Build/Best-Value", "Design-Build/Low-Bid", "Progressive Design-Build"},
    },
    {
        "id": "R7",
        "name": "High procurement cost limits Design-Build",
        "trigger": "E4 = A",
        "description": "Procurement cost would significantly limit competition. Design-Build has higher procurement costs including stipends for proposers.",
        "blocks": {"Design-Build/Best-Value"},
    },
    {
        "id": "R8",
        "name": "No in-house design expertise blocks DBB",
        "trigger": "F2 = C",
        "description": "Specialized design expertise not available in-house. DBB requires the Department to own and complete the full design. DB and PDB transfer design ownership to the contractor.",
        "blocks": {"Design-Bid-Build"},
    },
    {
        "id": "R9",
        "name": "Inadequate procurement expertise blocks complex methods",
        "trigger": "F1 = A",
        "description": "Department lacks resources or expertise for complex procurement. Design-Build and PDB require sophisticated two-phase procurement (RFQ + RFP).",
        "blocks": {"Design-Build/Best-Value", "Design-Build/Low-Bid", "Progressive Design-Build"},
    },
]


def _apply_overrides(recommended: str, runner_up: str, rating_lookup: dict) -> tuple:
    """Apply document-based override rules that block or favor certain methods.

    Returns: (recommended, runner_up, list_of_override_reason_strings)
    """
    blocked = set()
    reasons = []
    favor = None

    # Rule 1: A1=C — early stage blocks DBB, Design-Seq
    if rating_lookup.get("A1") == "C":
        blocked.update({"Design-Bid-Build", "Design-Sequencing"})
        reasons.append(OVERRIDE_RULES[0]["description"])

    # Rule 2: A2=C — large project blocks DBB
    if rating_lookup.get("A2") == "C":
        blocked.add("Design-Bid-Build")
        reasons.append(OVERRIDE_RULES[1]["description"])

    # Rule 3: A3=C — very complex blocks DBB, Design-Seq
    if rating_lookup.get("A3") == "C":
        blocked.update({"Design-Bid-Build", "Design-Sequencing"})
        reasons.append(OVERRIDE_RULES[2]["description"])

    # Rule 4: B1=C AND B2=C — max schedule compression blocks DBB
    if rating_lookup.get("B1") == "C" and rating_lookup.get("B2") == "C":
        blocked.add("Design-Bid-Build")
        reasons.append(OVERRIDE_RULES[3]["description"])

    # Rule 5: C2=C — performance specs favor DB/Best-Value
    if rating_lookup.get("C2") == "C":
        favor = "Design-Build/Best-Value"
        reasons.append(OVERRIDE_RULES[4]["description"])

    # Rule 6: E3=A — limited funding blocks DB, PDB
    if rating_lookup.get("E3") == "A":
        blocked.update({"Design-Build/Best-Value", "Design-Build/Low-Bid", "Progressive Design-Build"})
        reasons.append(OVERRIDE_RULES[5]["description"])

    # Rule 7: E4=A — high procurement cost blocks DB/Best-Value
    if rating_lookup.get("E4") == "A":
        blocked.add("Design-Build/Best-Value")
        reasons.append(OVERRIDE_RULES[6]["description"])

    # Rule 8: F2=C — no design resources blocks DBB
    if rating_lookup.get("F2") == "C":
        blocked.add("Design-Bid-Build")
        reasons.append(OVERRIDE_RULES[7]["description"])

    # Rule 9: F1=A — inadequate procurement blocks DB, PDB
    if rating_lookup.get("F1") == "A":
        blocked.update({"Design-Build/Best-Value", "Design-Build/Low-Bid", "Progressive Design-Build"})
        reasons.append(OVERRIDE_RULES[8]["description"])

    # Apply favor rule (only if not blocked)
    if favor and favor not in blocked and recommended != favor:
        runner_up = recommended
        recommended = favor

    # If recommended is blocked, find best non-blocked alternative
    if recommended in blocked:
        # Try runner_up first
        if runner_up and runner_up not in blocked:
            recommended, runner_up = runner_up, recommended
        else:
            # Walk fallback hierarchy
            for method in _FALLBACK_ORDER:
                if method not in blocked and method != recommended:
                    old = recommended
                    recommended = method
                    runner_up = old
                    break

    # If runner_up is also blocked, pick next available
    if runner_up in blocked:
        for method in _FALLBACK_ORDER:
            if method not in blocked and method != recommended:
                runner_up = method
                break

    return recommended, runner_up, reasons


def _compute_override_status(rating_lookup: dict) -> list:
    """Evaluate all 9 override rules and return each with triggered status."""
    statuses = []
    checks = [
        ("R1", lambda rl: rl.get("A1") == "C"),
        ("R2", lambda rl: rl.get("A2") == "C"),
        ("R3", lambda rl: rl.get("A3") == "C"),
        ("R4", lambda rl: rl.get("B1") == "C" and rl.get("B2") == "C"),
        ("R5", lambda rl: rl.get("C2") == "C"),
        ("R6", lambda rl: rl.get("E3") == "A"),
        ("R7", lambda rl: rl.get("E4") == "A"),
        ("R8", lambda rl: rl.get("F2") == "C"),
        ("R9", lambda rl: rl.get("F1") == "A"),
    ]
    for rule in OVERRIDE_RULES:
        rid = rule["id"]
        check_fn = next((fn for r_id, fn in checks if r_id == rid), None)
        triggered = check_fn(rating_lookup) if check_fn else False
        statuses.append({
            "rule_id": rid,
            "rule_name": rule["name"],
            "trigger_condition": rule["trigger"],
            "description": rule["description"],
            "triggered": triggered,
            "blocks": list(rule.get("blocks", set())),
            "favors": rule.get("favor", ""),
        })
    return statuses


def _build_comparison(recommended: str, runner_up: str, composite: float, section_scores: dict) -> str:
    """Build a qualitative comparison for borderline cases."""
    lines = [
        f"**Score: {composite:.2f} / 3.00**",
        "",
        f"The scores place this project near the boundary between **{recommended}** and **{runner_up}**.",
        "",
        "**Section Scores:**",
    ]
    section_names = {
        "A": "Project Scope & Characteristics",
        "B": "Schedule Issues",
        "C": "Opportunity for Innovation",
        "D": "Quality Enhancement",
        "E": "Cost Issues",
        "F": "Staffing Issues",
    }
    for s, name in section_names.items():
        score = section_scores.get(s, 2.0)
        lines.append(f"- {name}: {score:.2f} / 3.00 (weight: {SECTION_WEIGHTS[s]:.0%})")

    lines.extend([
        "",
        f"**{recommended}** is recommended as the primary method, but the project team should also evaluate "
        f"**{runner_up}** given the proximity of scores. Consider the specific project constraints, district "
        f"experience with each method, and stakeholder preferences when making the final decision.",
    ])
    return "\n".join(lines)


# ==============================================================================
# EXCEL EXPORT
# ==============================================================================

# Shared styles for Excel generation
def _get_styles():
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    styles = {}
    styles['hdr_fill'] = PatternFill("solid", fgColor="1F4E79")
    styles['hdr_font'] = Font(bold=True, color="FFFFFF", size=11)
    styles['even_fill'] = PatternFill("solid", fgColor="EBF3FB")
    styles['bdr'] = Border(
        left=Side("thin", "000000"), right=Side("thin", "000000"),
        top=Side("thin", "000000"), bottom=Side("thin", "000000"),
    )
    styles['wrap'] = Alignment(wrap_text=True, vertical="top")
    styles['center'] = Alignment(horizontal="center", vertical="center", wrap_text=True)
    return styles

def _title(ws, text, cols):
    from openpyxl.styles import Font, Alignment
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=cols)
    c = ws.cell(row=1, column=1, value=text)
    c.font = Font(bold=True, size=14, color="1F4E79")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

def _header_row(ws, row, headers):
    s = _get_styles()
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=ci, value=h)
        c.font, c.fill, c.border, c.alignment = s['hdr_font'], s['hdr_fill'], s['bdr'], s['center']

def _data_row(ws, row, values):
    s = _get_styles()
    fill = s['even_fill'] if row % 2 == 0 else None
    for ci, val in enumerate(values, 1):
        c = ws.cell(row=row, column=ci, value=val)
        c.border, c.alignment = s['bdr'], s['wrap']
        if fill:
            c.fill = fill
    return fill

def _used_bounds(ws):
    max_r = 0
    max_c = 0
    for r in range(1, ws.max_row + 1):
        row_has_data = False
        for c in range(1, ws.max_column + 1):
            v = ws.cell(row=r, column=c).value
            if v is not None and str(v).strip() != "":
                row_has_data = True
                max_c = max(max_c, c)
        if row_has_data:
            max_r = r
    return max_r, max_c

def _apply_template_design(ws):
    from openpyxl.styles import Alignment
    s = _get_styles()
    max_r, max_c = _used_bounds(ws)
    if max_r == 0 or max_c == 0:
        return

    ws.freeze_panes = "A2"
    ws.sheet_view.zoomScale = 90

    # Base styling on all used cells
    for r in range(1, max_r + 1):
        ws.row_dimensions[r].height = 20
        non_empty = 0
        row_texts = []
        for c in range(1, max_c + 1):
            cell = ws.cell(row=r, column=c)
            val = cell.value
            txt = str(val).strip() if val is not None else ""
            if txt:
                non_empty += 1
                row_texts.append(txt.lower())
            cell.font = s['body_font']
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = s['bdr']

        row_blob = " ".join(row_texts)
        is_title = ("project delivery selection tool" in row_blob) or ("project summary worksheet" in row_blob)
        is_header = any(keyword in row_blob for keyword in ["question", "worksheet", "scoring summary", "final selection", "criteria"])
        
        if is_title:
            for c in range(1, max_c + 1):
                cell = ws.cell(row=r, column=c)
                cell.font = s['title_font']
                if cell.value is not None and str(cell.value).strip():
                    cell.fill = s['template_hdr_fill']
                    cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
            ws.row_dimensions[r].height = 24
        elif is_header:
            for c in range(1, max_c + 1):
                cell = ws.cell(row=r, column=c)
                if cell.value is not None and str(cell.value).strip():
                    cell.fill = s['template_subhdr_fill']
                    cell.font = s['subhdr_font']

def build_evaluation_excel(eval_data: dict, recommendation: dict, project_name: str,
                           multi_method_data: dict = None, validation_data: dict = None) -> BytesIO:
    """Build a styled 5-sheet Excel workbook with full analysis."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    center = Alignment(horizontal="center", vertical="center")
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    # ===== Sheet 1: Executive Dashboard =====
    ws1 = wb.create_sheet("Dashboard")
    _title(ws1, f"Project Delivery Evaluation — {project_name}", 4)

    # Summary block
    ws1.cell(row=3, column=1, value="Recommended Method:").font = Font(bold=True, size=12)
    ws1.cell(row=3, column=2, value=recommendation.get("recommended_method", "")).font = Font(bold=True, size=13, color="166534")
    ws1.cell(row=4, column=1, value="Runner-Up:").font = Font(bold=True)
    ws1.cell(row=4, column=2, value=recommendation.get("runner_up_method", ""))
    ws1.cell(row=5, column=1, value="Composite Score:").font = Font(bold=True)
    ws1.cell(row=5, column=2, value=f"{recommendation.get('composite_score', 0):.3f} / 3.000")
    ws1.cell(row=6, column=1, value="Borderline Case:").font = Font(bold=True)
    ws1.cell(row=6, column=2, value="⚠ YES" if recommendation.get("is_borderline") else "No")

    overrides_triggered = [o for o in recommendation.get("override_status", []) if o.get("triggered")]
    ws1.cell(row=7, column=1, value="Overrides Triggered:").font = Font(bold=True)
    ws1.cell(row=7, column=2, value=str(len(overrides_triggered)))

    # Section score table
    ws1.cell(row=9, column=1, value="Section Score Breakdown").font = Font(bold=True, size=12, color="1F4E79")
    _header_row(ws1, 10, ["Section", "Avg Score", "Weight", "Weighted"])
    section_names = {
        "A": "Project Scope & Characteristics", "B": "Schedule Issues",
        "C": "Opportunity for Innovation", "D": "Quality Enhancement",
        "E": "Cost Issues", "F": "Staffing Issues",
    }
    section_scores = recommendation.get("section_scores", {})
    weighted_scores = recommendation.get("weighted_scores", {})
    for ri, (sec, name) in enumerate(section_names.items(), 11):
        avg = section_scores.get(sec, 2.0)
        wt = SECTION_WEIGHTS.get(sec, 0)
        ws_val = weighted_scores.get(sec, avg * wt)
        _data_row(ws1, ri, [f"{sec}: {name}", f"{avg:.3f}", f"{wt:.0%}", f"{ws_val:.4f}"])

    # Key drivers
    key_drivers = recommendation.get("key_drivers", [])
    ws1.cell(row=18, column=1, value="Top 5 Key Drivers").font = Font(bold=True, size=12, color="1F4E79")
    _header_row(ws1, 19, ["Question", "Rating", "Section", "Weighted Contribution"])
    for ri, kd in enumerate(key_drivers, 20):
        _data_row(ws1, ri, [kd["question_id"], kd["rating"], kd["section"], f"{kd['weighted_contribution']:.4f}"])

    ws1.column_dimensions["A"].width = 38
    ws1.column_dimensions["B"].width = 22
    ws1.column_dimensions["C"].width = 14
    ws1.column_dimensions["D"].width = 22

    # ===== Sheet 2: Comprehensive Evaluation (Full Rubric) =====
    ws2 = wb.create_sheet("Comprehensive Evaluation")
    _title(ws2, "Full Rubric Matrix", 8)

    headers2 = ["Q#", "Question", "Rating", "Raw Score", "Evidence", "Confidence", "Missing?", "Section Weight"]
    _header_row(ws2, 2, headers2)

    ratings = eval_data.get("ratings", [])
    raw_scores = recommendation.get("raw_scores", {})
    for ri, r in enumerate(ratings, 3):
        qid = r.get("question_id", "")
        rating = r.get("selected_rating", "")
        raw = raw_scores.get(qid, RATING_VALUES.get(rating, 2))
        sec = qid[0] if qid else ""
        sw = SECTION_WEIGHTS.get(sec, 0)
        vals = [
            qid, r.get("question_text", ""), rating, raw,
            r.get("extracted_evidence", ""),
            round(r.get("confidence", 0), 2),
            "Yes" if r.get("missing_info") else "No",
            f"{sw:.0%}",
        ]
        _data_row(ws2, ri, vals)
        # Color-code the rating cell
        rc = ws2.cell(row=ri, column=3)
        rc.alignment = center
        if rating == "A":
            rc.font = Font(bold=True, color="166534")
        elif rating == "B":
            rc.font = Font(bold=True, color="854D0E")
        elif rating == "C":
            rc.font = Font(bold=True, color="991B1B")

    ws2.column_dimensions["A"].width = 6
    ws2.column_dimensions["B"].width = 50
    ws2.column_dimensions["C"].width = 8
    ws2.column_dimensions["D"].width = 11
    ws2.column_dimensions["E"].width = 60
    ws2.column_dimensions["F"].width = 12
    ws2.column_dimensions["G"].width = 10
    ws2.column_dimensions["H"].width = 14

    # ===== Sheet 3: Overrides & Constraints =====
    ws3 = wb.create_sheet("Overrides & Constraints")
    _title(ws3, "Override Rule Tracking", 6)

    headers3 = ["Rule ID", "Rule Name", "Trigger", "Status", "Blocks", "Description"]
    _header_row(ws3, 2, headers3)

    override_status = recommendation.get("override_status", [])
    for ri, o in enumerate(override_status, 3):
        triggered = o.get("triggered", False)
        status_text = "✅ TRIGGERED" if triggered else "—"
        blocks = ", ".join(o.get("blocks", [])) or o.get("favors", "—")
        vals = [
            o.get("rule_id", ""), o.get("rule_name", ""),
            o.get("trigger_condition", ""), status_text,
            blocks, o.get("description", ""),
        ]
        fill = _data_row(ws3, ri, vals)
        if triggered:
            for ci in range(1, 7):
                ws3.cell(row=ri, column=ci).font = Font(bold=True, color="991B1B")

    ws3.column_dimensions["A"].width = 10
    ws3.column_dimensions["B"].width = 40
    ws3.column_dimensions["C"].width = 20
    ws3.column_dimensions["D"].width = 16
    ws3.column_dimensions["E"].width = 45
    ws3.column_dimensions["F"].width = 70

    # ===== Sheet 4: Multi-Method Comparison =====
    if multi_method_data:
        ws4 = wb.create_sheet("Method Comparison")
        _title(ws4, "All Delivery Methods — Suitability Ranking", 7)

        headers4 = ["Rank", "Method", "Score", "Status", "Pros", "Cons", "Key Factors"]
        _header_row(ws4, 2, headers4)

        for ri, ms in enumerate(multi_method_data.get("method_scores", []), 3):
            status = "🚫 BLOCKED" if ms.get("blocked") else "✅ Eligible"
            block_detail = " | ".join(ms.get("block_reasons", []))
            if block_detail:
                status += f" ({block_detail})"
            vals = [
                ms.get("rank", ""), ms.get("method", ""),
                f"{ms.get('score', 0):.4f}", status,
                " • ".join(ms.get("pros", [])),
                " • ".join(ms.get("cons", [])),
                " | ".join(ms.get("key_factors", [])),
            ]
            _data_row(ws4, ri, vals)
            # Color code: green for #1, amber for #2-3, gray for blocked
            rank = ms.get("rank", 99)
            for ci in range(1, 8):
                cell = ws4.cell(row=ri, column=ci)
                if ms.get("blocked"):
                    cell.font = Font(color="6B7280", italic=True)
                elif rank == 1:
                    cell.font = Font(bold=True, color="166534")
                elif rank <= 3:
                    cell.font = Font(color="854D0E")

        # Borderline comparison row if applicable
        bc = multi_method_data.get("borderline_comparison")
        if bc and bc.get("is_close"):
            br = ri + 2 if multi_method_data.get("method_scores") else 4
            ws4.merge_cells(start_row=br, start_column=1, end_row=br, end_column=7)
            c = ws4.cell(row=br, column=1,
                         value=f"⚠ BORDERLINE: Top methods within {bc['score_gap']:.4f} — recommend detailed comparison")
            c.font = Font(bold=True, color="B45309", size=11)

        ws4.column_dimensions["A"].width = 6
        ws4.column_dimensions["B"].width = 28
        ws4.column_dimensions["C"].width = 10
        ws4.column_dimensions["D"].width = 30
        ws4.column_dimensions["E"].width = 55
        ws4.column_dimensions["F"].width = 55
        ws4.column_dimensions["G"].width = 40

    # ===== Sheet 5: Validation Report =====
    if validation_data:
        ws5 = wb.create_sheet("Validation Report")
        _title(ws5, "AI vs District Rating Validation", 8)

        # Summary block
        summary = validation_data.get("summary", {})
        ws5.cell(row=3, column=1, value="Agreement Rate:").font = Font(bold=True, size=12)
        rate = summary.get("agreement_rate", 0)
        rate_color = "166534" if rate >= 80 else ("B45309" if rate >= 60 else "991B1B")
        ws5.cell(row=3, column=2, value=f"{rate}%").font = Font(bold=True, size=13, color=rate_color)
        ws5.cell(row=4, column=1, value="Questions Compared:").font = Font(bold=True)
        ws5.cell(row=4, column=2, value=summary.get("total_compared", 0))
        ws5.cell(row=5, column=1, value="Matches:").font = Font(bold=True)
        ws5.cell(row=5, column=2, value=summary.get("matches", 0))
        ws5.cell(row=6, column=1, value="Minor Mismatches:").font = Font(bold=True)
        ws5.cell(row=6, column=2, value=summary.get("minor_mismatches", 0))
        ws5.cell(row=7, column=1, value="Major Mismatches:").font = Font(bold=True)
        ws5.cell(row=7, column=2, value=summary.get("major_mismatches", 0))

        # Deviation impact
        dev = validation_data.get("deviation_impact", {})
        ws5.cell(row=9, column=1, value="Deviation Impact").font = Font(bold=True, size=12, color="1F4E79")
        ws5.cell(row=10, column=1, value="AI Recommendation:").font = Font(bold=True)
        ws5.cell(row=10, column=2, value=dev.get("ai_method", ""))
        ws5.cell(row=11, column=1, value="User-Based Recommendation:").font = Font(bold=True)
        ws5.cell(row=11, column=2, value=dev.get("user_method", ""))
        changed = dev.get("recommendation_changed", False)
        ws5.cell(row=12, column=1, value="Recommendation Changed?").font = Font(bold=True)
        chg_cell = ws5.cell(row=12, column=2, value="⚠ YES" if changed else "No")
        if changed:
            chg_cell.font = Font(bold=True, color="991B1B")

        # Detail table
        headers5 = ["Q#", "Question", "AI Rating", "User Rating", "Match", "Severity", "AI Evidence", "Confidence"]
        _header_row(ws5, 14, headers5)

        for ri, comp in enumerate(validation_data.get("comparisons", []), 15):
            severity = comp.get("severity", "match")
            if severity == "match":
                sev_display, sev_emoji = "Match", "✅"
            elif severity == "minor_mismatch":
                sev_display, sev_emoji = "Minor", "⚠"
            else:
                sev_display, sev_emoji = "Major", "🔴"
            vals = [
                comp.get("question_id", ""), comp.get("question_text", ""),
                comp.get("ai_rating", ""), comp.get("user_rating", ""),
                sev_emoji, sev_display,
                comp.get("ai_evidence", ""), comp.get("ai_confidence", 0),
            ]
            _data_row(ws5, ri, vals)
            # Color-code
            for ci in range(1, 9):
                cell = ws5.cell(row=ri, column=ci)
                if severity == "major_mismatch":
                    cell.font = Font(bold=True, color="991B1B")
                elif severity == "minor_mismatch":
                    cell.font = Font(color="B45309")

        ws5.column_dimensions["A"].width = 6
        ws5.column_dimensions["B"].width = 50
        ws5.column_dimensions["C"].width = 10
        ws5.column_dimensions["D"].width = 12
        ws5.column_dimensions["E"].width = 8
        ws5.column_dimensions["F"].width = 10
        ws5.column_dimensions["G"].width = 60
        ws5.column_dimensions["H"].width = 12

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _safe_sheet_title(title: str) -> str:
    """Return an Excel-safe sheet title (max 31 chars)."""
    invalid = ['\\', '/', '*', '?', ':', '[', ']']
    safe = title
    for ch in invalid:
        safe = safe.replace(ch, "-")
    return safe[:31]


def _apply_box_border(ws, start_row, start_col, end_row, end_col):
    """Utility to apply a thin border around a rectangular range of cells."""
    from openpyxl.styles import Border, Side
    thin = Side(border_style="thin", color="000000")
    for r in range(start_row, end_row + 1):
        for c in range(start_col, end_col + 1):
            cell = ws.cell(row=r, column=c)
            # Maintain existing border if any
            current = cell.border
            new_top = thin if r == start_row else current.top
            new_bottom = thin if r == end_row else current.bottom
            new_left = thin if c == start_col else current.left
            new_right = thin if c == end_col else current.right
            cell.border = Border(top=new_top, bottom=new_bottom, left=new_left, right=new_right)

def _populate_v2_summary_sheet(ws, q_list, rating_index, method_labels, project_name, multi_method_data, eval_data=None):
    """
    Specifically populates the 'Project Summary Worksheet' following the provided images.
    """
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    
    # Standard styles for this sheet
    thin = Side(border_style="thin", color="000000")
    bdr = Border(top=thin, left=thin, right=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    wrap = Alignment(wrap_text=True, vertical="top")
    blue_header = Font(bold=True, size=11, color="0000FF")
    # Grey-White Theme Colors
    grey_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    grey_bold = Font(bold=True, color="333333")

    # Column widths for "de-congested" look
    ws.column_dimensions['A'].width = 75
    ws.column_dimensions['B'].width = 25
    for ci in range(len(method_labels)):
        col_letter = get_column_letter(3 + ci*2 + 1)
        ws.column_dimensions[col_letter].width = 15.0

    # 1. INSTRUCTIONS
    ws.cell(row=1, column=1, value="INSTRUCTIONS").font = Font(bold=True)
    ws.cell(row=1, column=1).alignment = center
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=14)
    
    instructions = [
        "1. On the Project Summary Worksheet, complete the date of the review, project name, and selection committee members.",
        "2.  Answer all questions on Worksheet 1.  Record the score for each delivery method on the form as indicated.",
        "Note: if any one of the answers is \"No-Go,\" the delivery method need not be considered further for that project.",
        "3. After all the questions are answered, total the score for each delivery system and transfer the totals to the Scoring Summary section on the Project Summary Worksheet.",
        "4. Repeat steps 2 and 3 for Worksheet 2.",
        "5. Total the scores from Worksheets 1 and 2 the in Scoring Summary section of the Project Summary Worksheet.",
        "6. Select the project delivery method with the highest score and record any important selection committee comments in the space provided.",
        "Note: Complete one project delivery selection questionnaire for each unique project. If multiple project alternatives or subprojects are being considered, complete one questionnaire for each unique variation."
    ]
    for i, text in enumerate(instructions, 2):
        cell = ws.cell(row=i, column=1, value=text)
        cell.font = Font(size=9)
        ws.merge_cells(start_row=i, start_column=1, end_row=i, end_column=14)
    
    # Box the instructions
    # Style definitions to prevent NameErrors
    thin = Side(border_style="thin", color="000000")
    bold_font = Font(bold=True)
    grey_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center")
    
    curr = 12
    # 2. Project Headers
    tool_title = ws.cell(row=curr, column=4, value="Project Delivery Selection Tool")
    tool_title.font = Font(bold=True, size=12)
    tool_title.alignment = center
    ws.merge_cells(start_row=curr, start_column=4, end_row=curr, end_column=10)
    
    ws.cell(row=curr+1, column=4, value="Project Summary Worksheet").font = Font(bold=True, size=11)
    ws.cell(row=curr+1, column=4).alignment = center
    ws.merge_cells(start_row=curr+1, start_column=4, end_row=curr+1, end_column=10)
    
    # 2. Project Headers (Identity Table)
    curr += 1
    # Define identity rows
    district = eval_data.get("district", "N/A") if eval_data else "N/A"
    ea_num = eval_data.get("project_ea", "N/A") if eval_data else "N/A"
    
    identity = [
        ("Project Name:", project_name if project_name else "N/A"),
        ("Project District:", district),
        ("Project EA:", ea_num),
        ("Date of Review:", datetime.date.today().strftime("%m/%d/%y")),
    ]
    
    id_start_row = curr
    for label, val in identity:
        # Label cell
        l_cell = ws.cell(row=curr, column=1, value=label)
        l_cell.font = bold_font
        l_cell.fill = grey_fill
        l_cell.alignment = Alignment(horizontal="right")
        
        # Value cell
        v_cell = ws.cell(row=curr, column=2, value=val)
        v_cell.alignment = center
        ws.merge_cells(start_row=curr, start_column=2, end_row=curr, end_column=4)
        curr += 1
    
    _apply_box_border(ws, id_start_row, 1, curr-1, 4)
    
    curr += 1
    ws.cell(row=curr, column=1, value="Review is based on AI evaluation of project documentation").font = Font(italic=True, size=9)
    curr += 2
    
    # Committee Section
    comm_start = curr
    ws.cell(row=curr, column=1, value="Selection Committee Members:").font = bold_font
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=4)
    curr += 1
    for i in range(5):
        ws.cell(row=curr, column=1, value=f"{i+1}.").alignment = Alignment(horizontal="right")
        ws.cell(row=curr, column=2).border = Border(bottom=thin)
        ws.merge_cells(start_row=curr, start_column=2, end_row=curr, end_column=4)
        curr += 1
    _apply_box_border(ws, comm_start, 1, curr-1, 4)
    
    curr += 2
    # 3. SCORING SUMMARY
    sum_hdr = ws.cell(row=curr, column=1, value="SCORING SUMMARY")
    sum_hdr.font = Font(bold=True, size=11)
    sum_hdr.alignment = center
    sum_hdr.fill = header_fill
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=14)
    curr += 1
    
    # Method Headers for Summary Table
    table_start_row = curr
    ws.cell(row=curr, column=1, value="EVALUATION FACTORS").font = bold_font
    ws.cell(row=curr, column=1).alignment = center
    ws.cell(row=curr, column=1).fill = grey_fill
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=2)

    for ci, method in enumerate(method_labels):
        col = 3 + ci*2
        # Merge two columns for method label
        h_cell = ws.cell(row=curr, column=col, value=method)
        h_cell.font = bold_font
        h_cell.fill = grey_fill
        h_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.merge_cells(start_row=curr, start_column=col, end_row=curr, end_column=col+1)
    
    # Question column must be wide
    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 60
    
    # Increase row height for project headers to prevent cramping
    for r in range(13, 20):
        ws.row_dimensions[r].height = 25
    
    curr += 1
    
    # Pre-calculate scores per worksheet
    ws1_scores = {m: 0 for m in method_labels}
    ws2_scores = {m: 0 for m in method_labels}
    
    for q in q_list:
        qid = q["id"]
        sec = qid[0]
        sel = rating_index.get(qid, {}).get("selected_rating", "B").upper()
        for m in method_labels:
            pts = METHOD_AFFINITY.get(sec, {}).get(m, {}).get(sel, 0.5) * 10
            if sec == "A": ws1_scores[m] += pts
            else: ws2_scores[m] += pts

    summary_rows = [
        ("Project Scope and Characteristic Score (Worksheet 1)", ws1_scores),
        ("Success Criteria Score (Worksheet 2)", ws2_scores),
        ("Total Score", None) # Handled separately for styling
    ]
    
    for label, score_dict in summary_rows:
        label_cell = ws.cell(row=curr, column=1, value=label)
        label_cell.font = bold_font if score_dict else Font(bold=True, color="000000")
        label_cell.alignment = Alignment(wrap_text=True, horizontal="left", vertical="center")
        ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=2)
        
        for ci, m in enumerate(method_labels):
            col = 3 + ci*2
            val = round(score_dict[m]) if score_dict else round(ws1_scores[m] + ws2_scores[m])
            c = ws.cell(row=curr, column=col, value=val)
            c.font = bold_font
            c.alignment = center
            if not score_dict: # Total row
                c.fill = grey_fill
            ws.merge_cells(start_row=curr, start_column=col, end_row=curr, end_column=col+1)
        curr += 1
    
    # Apply borders to summary table
    _apply_box_border(ws, table_start_row, 1, curr-1, 3 + len(method_labels)*2 - 1)
    
    curr += 1
    ws.cell(row=curr, column=1, value="Final Selection:").font = bold_font
    curr += 1
    for ci, m in enumerate(method_labels):
        col = 1 + ci*2
        ws.cell(row=curr, column=col, value=f"☐ {m}").font = Font(size=9)
    curr += 2
    
    ws.cell(row=curr, column=1, value="Comments:").font = bold_font
    for _ in range(4):
        ws.cell(row=curr, column=2).border = Border(bottom=thin)
        ws.merge_cells(start_row=curr, start_column=2, end_row=curr, end_column=14)
        curr += 1

    curr += 2
    _v2_draw_questionnaire(ws, curr, q_list, rating_index, method_labels, ws1_scores, ws2_scores)


def _v2_draw_questionnaire(ws, start_row, q_list, rating_index, method_labels, ws1_scores, ws2_scores):
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    thin = Side(border_style="thin", color="000000")
    bdr = Border(top=thin, left=thin, right=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    bold_font = Font(bold=True)
    
    # Theme Colors
    grey_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    wrap = Alignment(wrap_text=True, vertical="top")

    curr = start_row
    
    # Worksheet 1
    ws.cell(row=curr, column=1, value="WORKSHEET 1").font = bold_font
    ws.cell(row=curr, column=1).alignment = center
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=14)
    curr += 1
    ws.cell(row=curr, column=1, value="EVALUATION OF PROJECT SCOPE AND CHARACTERISTICS").font = bold_font
    ws.cell(row=curr, column=1).alignment = center
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=14)
    curr += 2

    # Section A
    s_cell = ws.cell(row=curr, column=1, value="Project Scope and Characteristic Criteria")
    s_cell.font = bold_font
    s_cell.fill = header_fill
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=2)
    
    # Method Headers
    for ci, method in enumerate(method_labels):
        col = 3 + ci*2
        c = ws.cell(row=curr, column=col, value=method)
        c.font = Font(bold=True, size=9)
        c.fill = grey_fill
        c.alignment = Alignment(wrap_text=True, horizontal="center")
        ws.merge_cells(start_row=curr, start_column=col, end_row=curr, end_column=col+1)
    
    _apply_box_border(ws, curr, 1, curr, 3 + len(method_labels)*2 - 1)
    curr += 2
    
    a_questions = [q for q in q_list if q['id'].startswith("A")]
    for q in a_questions:
        qid = q["id"]
        sel_rating = rating_index.get(qid, {}).get("selected_rating", "B").upper()
        
        start_r = curr
        # Draw Question
        ws.cell(row=curr, column=1, value=f"{qid}. {q['question']}").font = bold_font
        ws.cell(row=curr, column=1).alignment = wrap
        curr += 1
        
        # Draw Options
        opts_rendered = 0
        for opt_key, opt_label in [("A", "option_a"), ("B", "option_b"), ("C", "option_c")]:
            opt_text = q.get(opt_label, "")
            if opt_text:
                ws.cell(row=curr, column=1, value=f"☐ {opt_key}. {opt_text}").font = Font(size=8)
                curr += 1
                opts_rendered += 1
        
        # Total rows for this question (Question + Options)
        end_r = curr - 1
        
        for ci, method in enumerate(method_labels):
            col = 3 + ci*2
            # Calculate points
            affinity = METHOD_AFFINITY.get("A", {}).get(method, {}).get(sel_rating, 0.5)
            pts = round(affinity * 5, 1)
            display_val = f"{sel_rating} ({pts})" if sel_rating else ""
            
            # ID cell (Merge vertically, leave empty)
            if start_r != end_r:
                ws.merge_cells(start_row=start_r, start_column=col, end_row=end_r, end_column=col)
            ws.cell(row=start_r, column=col).border = Border(left=thin, top=thin, bottom=thin)
            
            # Box cell (Merge vertically, show rating + points)
            if start_r != end_r:
                ws.merge_cells(start_row=start_r, start_column=col+1, end_row=end_r, end_column=col+1)
            c = ws.cell(row=start_r, column=col+1, value=display_val)
            c.font = bold_font
            c.border = Border(right=thin, top=thin, bottom=thin)
            c.alignment = center
            
        curr += 1 # Spacer row

    # SCORE Row for WS1
    ws.cell(row=curr, column=1, value="Project Characteristics Subtotal (Total Questions A1-A10)").font = Font(italic=True, size=9)
    ws.cell(row=curr, column=2, value="SCORE").font = bold_font
    for ci, m in enumerate(method_labels):
        col = 3 + ci*2
        c = ws.cell(row=curr, column=col+1, value=round(ws1_scores[m]))
        c.font = bold_font
        c.alignment = center
        c.border = Border(bottom=thin)
    curr += 4

    # Worksheet 2
    ws.cell(row=curr, column=1, value="WORKSHEET 2").font = bold_font
    ws.cell(row=curr, column=1).alignment = center
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=14)
    curr += 1
    ws.cell(row=curr, column=1, value="EVALUATION OF SUCCESS CRITERIA").font = bold_font
    ws.cell(row=curr, column=1).alignment = center
    ws.merge_cells(start_row=curr, start_column=1, end_row=curr, end_column=14)
    curr += 2

    other_sections = [
        ("B", "Schedule Issues"),
        ("C", "Opportunity for Innovation"),
        ("D", "Quality Enhancement"),
        ("E", "Cost Issues"),
        ("F", "Staffing Issues")
    ]
    
    for prefix, section_title in other_sections:
        ws.cell(row=curr, column=1, value=f"{prefix} - {section_title}").font = bold_font
        # Headers
        for ci, method in enumerate(method_labels):
            col = 3 + ci*2
            c = ws.cell(row=curr, column=col, value=method)
            c.font = Font(bold=True, size=8)
            c.alignment = center
            ws.merge_cells(start_row=curr, start_column=col, end_row=curr, end_column=col+1)
        curr += 1
        
        sec_questions = [q for q in q_list if q['id'].startswith(prefix)]
        for q in sec_questions:
            qid = q["id"]
            sel_rating = rating_index.get(qid, {}).get("selected_rating", "B").upper()
            
            start_r = curr
            # Question text
            ws.cell(row=curr, column=1, value=f"{qid}. {q['question']}").font = bold_font
            ws.cell(row=curr, column=1).alignment = wrap
            curr += 1
            
            # Options
            for opt_key, opt_label in [("A", "option_a"), ("B", "option_b"), ("C", "option_c")]:
                opt_text = q.get(opt_label, "")
                if opt_text:
                    ws.cell(row=curr, column=1, value=f"☐ {opt_key}. {opt_text}").font = Font(size=8)
                    curr += 1
            
            # Total rows for this question
            end_r = curr - 1
            
            for ci, method in enumerate(method_labels):
                col = 3 + ci*2
                # Calculate points
                affinity = METHOD_AFFINITY.get(prefix, {}).get(method, {}).get(sel_rating, 0.5)
                pts = round(affinity * 5, 1)
                display_val = f"{sel_rating} ({pts})" if sel_rating else ""
                
                # ID cell (Merge vertically, leave empty)
                if start_r != end_r:
                    ws.merge_cells(start_row=start_r, start_column=col, end_row=end_r, end_column=col)
                ws.cell(row=start_r, column=col).border = Border(left=thin, top=thin, bottom=thin)
                
                # Box cell (Merge vertically, show rating + points)
                if start_r != end_r:
                    ws.merge_cells(start_row=start_r, start_column=col+1, end_row=end_r, end_column=col+1)
                c = ws.cell(row=start_r, column=col+1, value=display_val)
                c.font = bold_font
                c.border = Border(right=thin, top=thin, bottom=thin)
                c.alignment = center
            
            curr += 1 # Spacer row

    # SCORE Row for WS2
    ws.cell(row=curr, column=1, value="Success Criteria Subtotal (Total questions B-F)").font = Font(italic=True, size=9)
    ws.cell(row=curr, column=2, value="SCORE").font = bold_font
    for ci, m in enumerate(method_labels):
        col = 3 + ci*2
        c = ws.cell(row=curr, column=col+1, value=round(ws2_scores[m]))
        c.font = bold_font
        c.alignment = center
        c.border = Border(bottom=thin)
    curr += 2
    
    ws.cell(row=curr, column=2, value="TOTAL").font = Font(bold=True)
    for ci, m in enumerate(method_labels):
        col = 3 + ci*2
        c = ws.cell(row=curr, column=col+1, value=round(ws1_scores[m] + ws2_scores[m]))
        c.font = Font(bold=True)
        c.fill = header_fill
        c.alignment = center
        c.border = bdr
    curr += 2

def _populate_rubric_sheet(ws, q_list, rating_index, method_labels=None, single_method=None, title=None, project_name=None):
    """
    Overhauled individual method worksheet with 7-column layout:
    ID, Criteria, Rating, Confidence, Source, Missing Info, effect.
    """
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    
    # Styles
    box_border = Border(top=Side(style='thin'), left=Side(style='thin'), right=Side(style='thin'), bottom=Side(style='thin'))
    thin_border = Side(style='thin')
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    top_left_align = Alignment(horizontal="left", vertical="top", wrap_text=True)
    bold_font = Font(bold=True)
    
    # Theme Colors
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    grey_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    wrap = Alignment(wrap_text=True, vertical="top")

    # 1. Header & Setup
    ws.column_dimensions['A'].width = 6   # ID
    ws.column_dimensions['B'].width = 50  # Criteria
    ws.column_dimensions['C'].width = 10  # Rating
    ws.column_dimensions['D'].width = 10  # Points
    ws.column_dimensions['E'].width = 10  # Confid.
    ws.column_dimensions['F'].width = 35  # Source Reasoning
    ws.column_dimensions['G'].width = 30  # Missing Info
    ws.column_dimensions['H'].width = 35  # Effect on Method

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
    title_cell = ws.cell(row=1, column=1, value=title if title else f"DETAILED EVALUATION: {single_method}")
    title_cell.font = Font(bold=True, size=14)
    title_cell.alignment = center_align

    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=8)
    sub_cell = ws.cell(row=2, column=1, value=f"Project: {project_name}" if project_name else "")
    sub_cell.alignment = center_align

    # 2. Table Headers
    headers = ["ID", "EVALUATION CRITERIA", "RATING", "POINTS", "CONFID.", "SOURCE REASONING", "MISSING INFO", "EFFECT ON METHOD"]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=4, column=ci, value=h)
        c.font = bold_font
        c.fill = header_fill
        c.alignment = center_align
        c.border = box_border
    
    current_row = 5
    for q in q_list:
        qid = q["id"]
        sec = qid[0]
        robj = rating_index.get(qid, {})
        sel_rating = robj.get("selected_rating", "").upper()
        confidence = robj.get("confidence", 0.0)
        
        # Extract new reasoning fields
        source_res = robj.get("source_reasoning", robj.get("extracted_evidence", "No evidence found"))
        missing_res = robj.get("missing_info_reasoning", "None")
        effect_res = robj.get("effect_on_method", robj.get("rubric_analysis", "Analysis N/A"))
        
        start_row = current_row
        
        # ID and Criteria
        ws.cell(row=current_row, column=1, value=qid).font = bold_font
        ws.cell(row=current_row, column=2, value=q["question"]).font = bold_font
        current_row += 1
        
        # Options
        for opt_key, opt_label in [("A", "option_a"), ("B", "option_b"), ("C", "option_c")]:
            opt_text = q.get(opt_label, "")
            if not opt_text: continue
            display_text = f"☐ {opt_key}. {opt_text}"
            if single_method:
                fit = METHOD_AFFINITY.get(sec, {}).get(single_method, {}).get(opt_key, 0.5)
                display_text += f" [{fit * 10.0:.1f} pts]"
        end_row = current_row - 1
        
        # Vertical Merging for ID and reasoning columns
        for col in [1, 3, 4, 5, 6, 7, 8]:
            if start_row != end_row:
                ws.merge_cells(start_row=start_row, start_column=col, end_row=end_row, end_column=col)
        
        # Rating Cell (C)
        affinity = METHOD_AFFINITY.get(sec, {}).get(single_method, {}).get(sel_rating, 0.5)
        pts = round(affinity * 5, 1)
        
        c_rate = ws.cell(row=start_row, column=3, value=sel_rating)
        c_rate.font = bold_font
        c_rate.alignment = center_align
        c_rate.border = box_border
        
        # Points Cell (D)
        c_pts = ws.cell(row=start_row, column=4, value=pts)
        c_pts.font = bold_font
        c_pts.alignment = center_align
        c_pts.border = box_border
        
        # Confidence Cell (E)
        c_conf = ws.cell(row=start_row, column=5, value=f"{confidence:.2f}")
        c_conf.alignment = center_align
        c_conf.border = box_border
        
        # Source Reasoning (F)
        c_src = ws.cell(row=start_row, column=6, value=source_res)
        c_src.alignment = top_left_align
        c_src.border = box_border
        c_src.font = Font(size=9)
        
        # Missing Info (G)
        c_miss = ws.cell(row=start_row, column=7, value=missing_res)
        c_miss.alignment = top_left_align
        c_miss.border = box_border
        c_miss.font = Font(size=9)
        
        # Effect on Method (H)
        c_eff = ws.cell(row=start_row, column=8, value=effect_res)
        c_eff.alignment = top_left_align
        c_eff.border = box_border
        c_eff.font = Font(size=9)
        
        # Borders for Criteria column (B)
        for r in range(start_row, end_row + 1):
            ws.cell(row=r, column=2).border = Border(left=thin_border, right=thin_border, top=thin_border if r == start_row else None, bottom=thin_border if r == end_row else None)

        current_row += 1 # Spacer
    ws.column_dimensions["L"].width = 10
    ws.column_dimensions["M"].width = 12

    ws.column_dimensions["B"].width = 100


def build_evaluation_excel_v2(
    eval_data: dict,
    recommendation: dict,
    project_name: str,
    template_path: str,
    multi_method_data: dict = None,
) -> BytesIO:
    """
    Build V2 workbook:
    - Keep the provided template sheet(s) as the summary presentation layer
    - Add one detailed sheet per delivery method
    """
    import openpyxl
    import pandas as pd
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    try:
        if str(template_path).lower().endswith(".xls") and not str(template_path).lower().endswith(".xlsx"):
            # Convert legacy .xls template into an in-memory openpyxl workbook via calamine.
            xls = pd.ExcelFile(template_path, engine="calamine")
            wb = openpyxl.Workbook()
            wb.remove(wb.active)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(template_path, sheet_name=sheet_name, header=None, engine="calamine")
                ws = wb.create_sheet(_safe_sheet_title(sheet_name))
                for r_idx in range(df.shape[0]):
                    for c_idx in range(df.shape[1]):
                        val = df.iat[r_idx, c_idx]
                        if pd.isna(val):
                            val = None
                        ws.cell(row=r_idx + 1, column=c_idx + 1, value=val)
        else:
            if os.path.exists(template_path):
                wb = openpyxl.load_workbook(template_path)
            else:
                wb = openpyxl.Workbook()
                # If creating fresh, remove default sheet so our added ones are at the start
                if "Sheet" in wb.sheetnames:
                    wb.remove(wb["Sheet"])
    except Exception as e:
        print(f"DEBUG: V2 Template load failed: {e}. Falling back to fresh workbook.")
        wb = openpyxl.Workbook()
        if "Sheet" in wb.sheetnames:
            wb.remove(wb["Sheet"])

    # Remove all existing template sheets to ensure the new summary is the single source of truth
    for ws in list(wb.worksheets):
        wb.remove(ws)

    import openpyxl

    ratings = eval_data.get("ratings", [])
    rating_index = {r.get("question_id"): r for r in ratings}
    
    try:
        if str(template_path).lower().endswith(".xls") and not str(template_path).lower().endswith(".xlsx"):
            import pandas as pd
            xls = pd.ExcelFile(template_path, engine="calamine")
            wb = openpyxl.Workbook()
            wb.remove(wb.active)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(template_path, sheet_name=sheet_name, header=None, engine="calamine")
                ws = wb.create_sheet(_safe_sheet_title(sheet_name))
                for r_idx in range(df.shape[0]):
                    for c_idx in range(df.shape[1]):
                        val = df.iat[r_idx, c_idx]
                        ws.cell(row=r_idx+1, column=c_idx+1, value=None if pd.isna(val) else val)
        else:
            wb = openpyxl.load_workbook(template_path) if os.path.exists(template_path) else openpyxl.Workbook()
    except:
        wb = openpyxl.Workbook()

    # Remove all existing sheets to rebuild strictly from scratch with our logic
    for ws in list(wb.worksheets):
        wb.remove(ws)

    method_labels = [
        "Design-Bid-Build", "Design-Sequencing", "Design-Build/Low-Bid",
        "Design-Build/Best-Value", "CM/GC", "Progressive Design-Build",
    ]

    # 1. Evaluation Summary (The main overview matching the template images)
    summary_ws = wb.create_sheet("Evaluation Summary")
    _populate_v2_summary_sheet(
        summary_ws, RUBRIC_QUESTIONS, rating_index, 
        method_labels=method_labels,
        project_name=project_name,
        multi_method_data=multi_method_data,
        eval_data=eval_data
    )
    
    # 2. Detailed Method Sheets (The 'elaborations' as extensions)
    for method in method_labels:
        # Create a safe sheet name (e.g. 'DB-BV Evaluation')
        safe_name = method.replace("Design-Build/", "DB-").split("/")[0]
        if len(safe_name) > 20: safe_name = safe_name[:20]
        m_ws = wb.create_sheet(f"{safe_name} Extension")
        
        _populate_rubric_sheet(
            m_ws, RUBRIC_QUESTIONS, rating_index,
            single_method=method,
            title=f"{method} Detailed Elaboration: {project_name}",
            project_name=project_name
        )

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

